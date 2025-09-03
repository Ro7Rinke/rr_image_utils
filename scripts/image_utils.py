import copy
import subprocess
import cv2
import fitz
import io
from multiprocessing import Pool, cpu_count
from debug_log import print_log
from session import get_session_images_path, new_uuid, copy_file, ensure_path
from concurrent.futures import ThreadPoolExecutor
from PIL import Image, UnidentifiedImageError
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
import numpy as np

Image.MAX_IMAGE_PIXELS = None  # sem limite

def import_image(image_info):
    src = image_info.get('external_source_path')
    dst = image_info.get('path')
    status, error = copy_file(src, dst)
    image_info['status'] = status
    image_info['error'] = error
    return image_info

def get_images_path_from_dirs(dir_images_path):
    new_images_path = []
    for image_path in dir_images_path:
        image_path = ensure_path(image_path)
        if image_path.is_dir():
            all_files = [p for p in image_path.iterdir() if p.is_file() and not p.name.startswith('.')]
            new_images_path.extend(all_files)
        else:
            new_images_path.append(image_path)
    
    return new_images_path

def import_images(session_id, images_path):
    images_folder_path = get_session_images_path(session_id)
    import_images_info = []
    images_info = []
    error_images_info = []

    images_path = get_images_path_from_dirs(images_path)

    for image_path in images_path:
        image_path = ensure_path(image_path)
        image_id = new_uuid()
        image_info = {
            'external_source_path': image_path,
            'id': image_id,
            'session_id': session_id,
            'path': images_folder_path / f'{image_id}--{image_path.name}'
        }
        import_images_info.append(image_info)
    
    with ThreadPoolExecutor() as executor:
        copy_results = list(executor.map(import_image, import_images_info))
        for result_image_info in copy_results:
            if result_image_info.get('status'):
                images_info.append(result_image_info)
            else:
                error_images_info.append(result_image_info)

    return images_info, error_images_info

def import_images_from_pdf(session_id, pdfs_path, page_as_image = False, dpi = 300):
    images_folder_path = get_session_images_path(session_id)
    images_info = []

    for pdf_path in pdfs_path:
        pdf_path = ensure_path(pdf_path)
        pdf_doc = fitz.open(pdf_path)

        for page_index in range(len(pdf_doc)):
            page = pdf_doc[page_index]
            page_number = page_index + 1

            if page_as_image:
                pix = page.get_pixmap(dpi=dpi)
                image_id = new_uuid()
                image_info = {
                    'external_source_path': images_folder_path / f'{pdf_path.stem}_{page_number}.png',
                    'id': image_id,
                    'session_id': session_id,
                    'path': images_folder_path / f'{image_id}--{pdf_path.stem}_{page_number}.png'
                }
                pix.save(image_info.get('path'))
                images_info.append(image_info)
            else:
                images = page.get_images(full=True)

                for image_index, image in enumerate(images):
                    xref = image[0]
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image['image']
                    image_ext = base_image['ext']
                    image_id = new_uuid()
                    image_info = {
                        'external_source_path': images_folder_path / f'{pdf_path.stem}_{page_number}.{image_ext}',
                        'id': image_id,
                        'session_id': session_id,
                        'path': images_folder_path / f'{image_id}--{pdf_path.stem}_{page_number}.{image_ext}'
                    }
                    with open(image_info.get('path'), 'wb') as file:
                        file.write(image_bytes)
                    
                    images_info.append(image_info)

    return images_info, []
    
def export_image(image_info):
    src = image_info.get('path')
    dst = image_info.get('external_output_path')
    status, error = copy_file(src, dst)
    image_info['status'] = status
    image_info['error'] = error
    return image_info

def export_images(images_info, output_directory_path, with_id = False, prefix = None, sufix = None):
    success_images_info = []
    error_images_info = []
    output_directory_path = ensure_path(output_directory_path)
    for image_info in images_info:
        if with_id:
            name = f"{image_info.get('external_source_path').stem}--{image_info.get('id')}{image_info.get('external_source_path').suffix}"
        else:
            name = image_info.get('external_source_path').name
        
        if prefix:
            name = f'{prefix}{name}'
        
        if sufix:
            name = f'{name}{sufix}'
            
        image_info['external_output_path'] = output_directory_path / name

    with ThreadPoolExecutor() as executor:
        copy_results = list(executor.map(export_image, images_info))
        for image_info in copy_results:
            if image_info.get('status'):
                success_images_info.append(image_info)
            else:
                error_images_info.append(image_info)

    return success_images_info, error_images_info

def save_new_image(image_info, new_img, format=None, dpi=None, quality=None, optimize=True):
    old_image_info = None
    images_folder_path = get_session_images_path(image_info.get('session_id'))
    new_image_info = copy.deepcopy(image_info)
    new_image_info['id'] = new_uuid()
    
    if image_info.get('id') is not None:
        new_image_info['old_id'] = image_info.get('id')
        old_image_info = image_info
    
    # decide extensão
    if format:
        ext = f".{format.lower()}"
    else:
        # usa a extensão que já estava salva (não a do external_source)
        ext = image_info.get('path').suffix
    
    filename = f'{image_info.get("external_source_path").stem}{ext}'
    id_filename = f"{new_image_info.get("id")}--{filename}"
    new_image_info['path'] = images_folder_path / id_filename
    
    # 🔑 atualiza também o "external_source_path" para refletir o novo formato
    new_image_info['external_source_path'] = images_folder_path / filename
    
    save_args = {'optimize': optimize}
    if format is not None:
        save_args['format'] = format
    if dpi is not None:
        save_args['dpi'] = (dpi, dpi)
    if quality is not None:
        save_args['quality'] = quality
    
    new_img.save(new_image_info.get('path'), **save_args)
    new_image_info['status'] = True

    return new_image_info, old_image_info

def resize_image(config):
    try:
        image_info = config.get('image_info')
        width = config.get('width')
        height = config.get('height')
        dpi = config.get('dpi')
        scale = config.get('scale')

        with Image.open(image_info.get('path')) as img:
            
            original_width, original_height = img.size

            if scale == 'percentage':
                width_px = original_width * (width / 100)
                height_px = original_height * (height / 100)
            else:
                match scale:
                    case 'px':
                        pts_divider = dpi
                    case 'mm':
                        pts_divider = 25.4
                    case 'cm':
                        pts_divider = 2.54
                    case 'm':
                        pts_divider = 0.0254
                    case 'in':
                        pts_divider = 1
                    case _:
                        pts_divider = 25.4

                width_px = width * dpi / pts_divider
                height_px = height * dpi / pts_divider

            if width_px > original_width and height_px > original_height:
                algorithm = Image.BICUBIC
            else:
                algorithm = Image.LANCZOS

            resized = img.resize((int(width_px), int(height_px)), algorithm)
            return save_new_image(image_info, resized)
        
    except FileNotFoundError:
        error = f"Arquivo não encontrado: {image_info.get('path')}"
    except UnidentifiedImageError:
        error = f"Arquivo não é uma imagem válida: {image_info.get('path')}"
    except OSError as e:
        error = f"Falha ao processar imagem '{image_info.get('path')}': {e}"
    except ValueError as e:
        error = f"Valor inválido ao salvar '{image_info.get('path')}': {e}"
    except KeyError as e:
        error = f"Formato não suportado para '{image_info.get('path')}': {e}"
    except Exception as e:
        error = f"Erro inesperado com '{image_info.get('path')}': {e}"
    
    image_info['status'] = False
    image_info['error'] = error

    return image_info, []

def resize_images(images_info, width = None, height = None, dpi = 300, scale = 'mm'):
    old_images_info = []
    new_images_info = []
    error_images_info = []
    configs = []

    for image_info in images_info:
        config = {
            'image_info': image_info,
            'width': width,
            'height': height,
            'dpi': dpi,
            'scale': scale
        }
        configs.append(config)

    with Pool(processes=cpu_count()) as pool:
        resize_results = pool.map(resize_image, configs)

    for new_image_info, old_image_info in resize_results:
        if new_image_info.get('status'):
            new_images_info.append(new_image_info)
            if isinstance(old_image_info, dict):
                old_images_info.append(old_image_info)
        else:
            error_images_info.append(new_image_info)

    return new_images_info, old_images_info, error_images_info

def convert_to_px(value, scale, dpi = 300, total_size = None):
    try:
        match(scale):
            case 'px':
                return int(value)
            case 'mm':
                return int(value * dpi / 25.4)
            case 'cm':
                return int(value * dpi / 2.54)
            case 'in':
                return int(value * dpi)
            case 'percentage':
                return int(total_size * value / 100)
            case _:
                raise ValueError(f'Escala não suportada: {scale}')
    except Exception as e:
        raise ValueError(f'Erro ao converter escala "{scale}": {e}')

def trim_transparent_borders(config):
    try:
        image_info = config.get('image_info')
        left = config.get('left')
        right = config.get('right')
        top = config.get('top')
        bottom = config.get('bottom')
        scale = config.get('scale')
        type = config.get('type')
        color = config.get('color')
        dpi = config.get('dpi')

        with Image.open(image_info.get('path')) as img:
            # Garante que tenha alpha para processar
            if img.mode in ("RGBA", "LA"):
                # Pega o canal alpha
                alpha = img.split()[-1]
                # bounding box dos pixels não transparentes
                bbox = alpha.getbbox()

                if bbox:
                    trimmed = img.crop(bbox)
                else:
                    # imagem é totalmente transparente → retorna ela mesmo
                    trimmed = img
            else:
                # Sem canal alpha → retorna ela mesmo
                trimmed = img

            return save_new_image(image_info, trimmed)

    except FileNotFoundError:
        error = f"Arquivo não encontrado: {image_info.get('path')}"
    except UnidentifiedImageError:
        error = f"Arquivo não é uma imagem válida: {image_info.get('path')}"
    except OSError as e:
        error = f"Falha ao processar imagem '{image_info.get('path')}': {e}"
    except Exception as e:
        error = f"Erro inesperado com '{image_info.get('path')}': {e}"

    image_info['status'] = False
    image_info['error'] = error

    return image_info, []

def remove_background_exact(config):
    try:
        image_info = config.get('image_info')
        left = config.get('left')
        right = config.get('right')
        top = config.get('top')
        bottom = config.get('bottom')
        scale = config.get('scale')
        type = config.get('type')
        color = config.get('color')
        dpi = config.get('dpi')

        with Image.open(image_info.get('path')).convert("RGBA") as img:
            datas = img.getdata()
            new_data = []
            for item in datas:
                # item é (R,G,B,A) ou (R,G,B)
                if item[0:3] == hex_to_rgb(color):  # compara só RGB
                    new_data.append((255, 255, 255, 0))  # transparente
                else:
                    new_data.append(item)
            img.putdata(new_data)

            return save_new_image(image_info, img, format='PNG')

    except FileNotFoundError:
        error = f"Arquivo não encontrado: {image_info.get('path')}"
    except UnidentifiedImageError:
        error = f"Arquivo não é uma imagem válida: {image_info.get('path')}"
    except OSError as e:
        error = f"Falha ao processar imagem '{image_info.get('path')}': {e}"
    except Exception as e:
        error = f"Erro inesperado com '{image_info.get('path')}': {e}"

    image_info['status'] = False
    image_info['error'] = error

    return image_info, []

def edit_border_image(config):
    image_info = config.get('image_info')
    left = config.get('left')
    right = config.get('right')
    top = config.get('top')
    bottom = config.get('bottom')
    scale = config.get('scale')
    type = config.get('type')
    color = config.get('color')
    dpi = config.get('dpi')

    try:
        with Image.open(image_info.get('path')) as img:
            original_width, original_height = img.size

            left_px = convert_to_px(left, scale, dpi, original_width)
            right_px = convert_to_px(right, scale, dpi, original_width)
            top_px = convert_to_px(top, scale, dpi, original_height)
            bottom_px = convert_to_px(bottom, scale, dpi, original_height)

            crop_box = (
                left_px,
                top_px,
                original_width - right_px,
                original_height - bottom_px
            )

            cropped_img = img.crop(crop_box)
            return save_new_image(image_info, cropped_img)

    except FileNotFoundError:
        error = f"Arquivo não encontrado: {image_info.get('path')}"
    except UnidentifiedImageError:
        error = f"Arquivo não é uma imagem válida: {image_info.get('path')}"
    except OSError as e:
        error = f"Falha ao processar imagem '{image_info.get('path')}': {e}"
    except ValueError as e:
        error = f"Valor inválido ao salvar '{image_info.get('path')}': {e}"
    except KeyError as e:
        error = f"Formato não suportado para '{image_info.get('path')}': {e}"
    except Exception as e:
        error = f"Erro inesperado com '{image_info.get('path')}': {e}"

    image_info['status'] = False
    image_info['error'] = error

    return image_info, []
    
def edit_border_images(images_info, left=0, right=0, top=0, bottom=0, scale='px', type = 'cut', color = None, dpi = 300):
    old_images_info = []
    new_images_info = []
    error_images_info = []
    configs = []

    for image_info in images_info:
        config = {
            'image_info': image_info,
            'left': left,
            'right': right,
            'top': top,
            'bottom': bottom,
            'scale': scale,
            'type': type,
            'color': color,
            'dpi': dpi
        }
        configs.append(config)

    with Pool(processes=cpu_count()) as pool:
        match(type):
            case 'cut':
                cropped_results = pool.map(edit_border_image, configs)
            case 'trim':
                cropped_results = pool.map(trim_transparent_borders, configs)
            case 'bg':
                cropped_results = pool.map(remove_background_exact, configs)

    for new_image_info, old_image_info in cropped_results:
        if new_image_info.get('status'):
            new_images_info.append(new_image_info)
            if isinstance(old_image_info, dict):
                old_images_info.append(old_image_info)
        else:
            error_images_info.append(new_image_info)

    return new_images_info, old_images_info, error_images_info

def export_to_word(images_info, output_directory_path, dpi = None, file_name = None):
    success_images_info = []
    error_images_info = []

    doc = Document()

    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.right_margin = Inches(1)
    section.left_margin = Inches(1)

    for image_info in images_info:
        error = None
        try:
            with Image.open(image_info.get('path')) as img:
                width_px, height_px = img.size
                image_dpi = dpi or img.info.get('dpi', (72, 72))[0]

                width_in = width_px / image_dpi
                height_in = height_px / image_dpi

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = paragraph.add_run()
                run.add_picture(str(image_info.get('path')), width=Inches(width_in), height=Inches(height_in))

                doc.add_page_break()

        except FileNotFoundError:
            error = f"Arquivo não encontrado: {image_info.get('path')}"
        except UnidentifiedImageError:
            error = f"Arquivo não é uma imagem válida: {image_info.get('path')}"
        except OSError as e:
            error = f"Falha ao processar imagem '{image_info.get('path')}': {e}"
        except ValueError as e:
            error = f"Valor inválido ao salvar '{image_info.get('path')}': {e}"
        except KeyError as e:
            error = f"Formato não suportado para '{image_info.get('path')}': {e}"
        except Exception as e:
            error = f"Erro inesperado com '{image_info.get('path')}': {e}"
        
        if error:
            image_info['error'] = error
            image_info['status'] = False
            error_images_info.append(image_info)
        else:
            image_info['status'] = True
            success_images_info.append(image_info)

    if doc.paragraphs[-1].text == '':
        doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

    output_directory_path = ensure_path(output_directory_path)
    file_name = f'doc_of_images.docx' if file_name is None else f'{file_name}.docx'
    output_path =  output_directory_path / file_name
    doc.save(output_path)

    return success_images_info, error_images_info

def rotate_if_needed(img):
    img_width, img_height = img.size
    if img_width > img_height:    
        return img.rotate(90, expand=True)
    return img

def create_pdf_page(config):
    image_info = config.get('image_info')
    dpi = config.get('dpi')

    try:
        with Image.open(image_info.get('path')) as img:

            if img.mode in ("RGBA", "LA"):
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[3])
                img = background
            else:
                img = img.convert("RGB")

            img = rotate_if_needed(img)

            img_width, img_height = img.size
            image_dpi = dpi or img.info.get('dpi', (72, 72))[0]

            A4_WIDTH_PX = int(8.27 * image_dpi)
            A4_HEIGHT_PX = int(11.69 * image_dpi)
            A4_SIZE = (A4_WIDTH_PX, A4_HEIGHT_PX)

            x = (A4_WIDTH_PX - img_width) // 2
            y = (A4_HEIGHT_PX - img_height) // 2

            page = Image.new("RGB", A4_SIZE, (255, 255, 255))
            page.paste(img, (x, y))
            
            return save_new_image(image_info, page)

    except FileNotFoundError:
        error = f"Arquivo não encontrado: {image_info.get('path')}"
    except UnidentifiedImageError:
        error = f"Arquivo não é uma imagem válida: {image_info.get('path')}"
    except OSError as e:
        error = f"Falha ao processar imagem '{image_info.get('path')}': {e}"
    except ValueError as e:
        error = f"Valor inválido ao salvar '{image_info.get('path')}': {e}"
    except KeyError as e:
        error = f"Formato não suportado para '{image_info.get('path')}': {e}"
    except Exception as e:
        error = f"Erro inesperado com '{image_info.get('path')}': {e}"

    image_info['status'] = False
    image_info['error'] = error

    return image_info, None

def export_to_pdf(images_info, output_directory_path, dpi=None, file_name = None):
    success_images_info = []
    error_images_info = []
    configs = []
    images = []
    
    file_name = f'pdf_of_images.pdf' if file_name is None else f'{file_name}.pdf'
    output_directory_path = ensure_path(output_directory_path)
    output_pdf_path =  output_directory_path / file_name
    
    # c = canvas.Canvas(str(output_pdf_path), pagesize=A4)

    for image_info in images_info:
        config = {
            'image_info': image_info,
            'dpi': dpi
        }
        configs.append(config)

    with Pool(processes=cpu_count()) as pool:
        pages_result = pool.map(create_pdf_page, configs)

    for image_info, old_image_info in pages_result:
        if image_info.get('status'):
            img = Image.open(image_info.get('path'))
            images.append(img)
            success_images_info.append(old_image_info)
        else:
            error_images_info.append(image_info)

    resolution = dpi or images[0].info.get('dpi', (72, 72))[0]
    if images:
        images[0].save(
            output_pdf_path,
            save_all=True,
            append_images=images[1:],
            resolution=resolution,
            optimize=True
        )

        for img in images:
            img.close()

    return success_images_info, error_images_info

def get_from_grid(config):
    image_info = config.get('image_info')
    rows = config.get('rows')
    cols = config.get('cols')
    new_images_info = []

    # try:
    with Image.open(image_info.get('path')) as img:
        img_width, img_height = img.size
        cell_width = img_width // cols
        cell_height = img_height // rows

        for row in range(rows):
            for col in range(cols):
                left = col * cell_width
                upper = row * cell_height
                right = left + cell_width
                lower = upper + cell_height

                box = (left, upper, right, lower)
                cropped = img.crop(box)

                original_name = image_info.get('path').stem
                new_image_info = {
                    'external_source_path': ensure_path(f"{original_name}_{row}_{col}.png"),
                    'path': ensure_path(f"{original_name}_{row}_{col}.png"),
                    'session_id': image_info.get('session_id')
                }
                new_image_info = save_new_image(new_image_info, cropped)
                new_images_info.append(new_image_info)
        
        return new_images_info

def images_from_grid(images_info, rows = 1, cols = 1):
    new_images_info = []
    configs = []

    for image_info in images_info:
        config = {
            'image_info': image_info,
            'cols': cols,
            'rows': rows
        }
        configs.append(config)

    with Pool(processes=cpu_count()) as pool:
        cropped_results = pool.map(get_from_grid, configs)

    new_images_info = [img[0] for sublist in cropped_results for img in sublist]

    return new_images_info

def quicklook_images(images_info):
    paths = []
    
    for image_info in images_info:
        image_path = ensure_path(image_info.get('path'))
        if image_path.is_file():
            paths.append(str(image_path))

    subprocess.run(["qlmanage", "-p"] + paths)

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        # duplicar cada caractere, ex: "F3A" -> "FF33AA"
        hex_color = ''.join(c*2 for c in hex_color)
    elif len(hex_color) != 6:
        raise ValueError(f"Cor hexadecimal inválida: {hex_color}")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def convert_to_jpeg(config):
    image_info = config.get('image_info')
    background_color = config.get('background_color')
    dpi = config.get('dpi')
    quality = config.get('quality')

    try:
        with Image.open(image_info.get('path')) as img:
            bg_rgb = hex_to_rgb(background_color)

            if img.mode in ("RGBA", "LA"):
                background = Image.new("RGB", img.size, bg_rgb)
                background.paste(img, mask=img.split()[-1])
                img = background
            elif img.mode != "RGB":
                img = img.convert("RGB")

            return save_new_image(image_info, img, format="JPEG", dpi=dpi, quality=quality)
            
    except FileNotFoundError:
        error = f"Arquivo não encontrado: {image_info.get('path')}"
    except UnidentifiedImageError:
        error = f"Arquivo não é uma imagem válida: {image_info.get('path')}"
    except OSError as e:
        error = f"Falha ao processar imagem '{image_info.get('path')}': {e}"
    except ValueError as e:
        error = f"Valor inválido ao salvar '{image_info.get('path')}': {e}"
    except KeyError as e:
        error = f"Formato não suportado para '{image_info.get('path')}': {e}"
    except Exception as e:
        error = f"Erro inesperado com '{image_info.get('path')}': {e}"

    image_info['status'] = False
    image_info['error'] = error

    return image_info

def convert_images_to_jpeg(images_info, dpi=None, quality=85, background_color='#FFFFFF'):
    new_images_info = []
    old_images_info = []
    error_images_info = []
    configs = []

    for image_info in images_info:
        config = {
            'image_info': image_info,
            'background_color': background_color,
            'dpi': dpi,
            'quality': quality
        }
        configs.append(config)

    with Pool(processes=cpu_count()) as pool:
        converted_results = pool.map(convert_to_jpeg, configs)

    for new_image_info, old_image_info in converted_results:
        if new_image_info.get('status'):
            new_images_info.append(new_image_info)
            if isinstance(old_image_info, dict):
                old_images_info.append(old_image_info)
        else:
            error_images_info.append(new_image_info)

    return new_images_info, old_images_info, error_images_info

def convert_cv2_to_pil(cv2_img):
    if len(cv2_img.shape) == 2:
        return Image.fromarray(cv2_img)
    elif len(cv2_img.shape) == 3:
        channels = cv2_img.shape[2]

        if channels == 4:
            img = cv2.cvtColor(cv2_img, cv2.COLOR_BGRA2RGBA)
            return Image.fromarray(img)
        elif channels == 3:
            img = cv2.cvtColor(cv2_img, cv2.COLOR_BGR2RGB)
            return Image.fromarray(img)
        
    raise ValueError("Formato de imagem desconhecido")

def remove_noise_from_image(image_info):

    try:
        img = cv2.imread(str(image_info.get('path')))
        if img is None:
            image_info['error'] = f"Imagem não pôde ser lida: {image_info.get('path')}"
            image_info['status'] = False
            
            return image_info, []
        
        # Configurações dos filtros
        GAUSSIAN_BLUR_ENABLED = True
        GAUSSIAN_BLUR_KERNEL_SIZE = (5, 5)  # Tamanho do kernel para o desfoque gaussiano
        GAUSSIAN_BLUR_SIGMA_X = 1.5 #1.2 #4.0 #1.0  # Desvio padrão no eixo X
        GAUSSIAN_BLUR_SIGMA_Y = 1.5 #1.2 #4.0 #1.0  # Desvio padrão no eixo Y

        BILATERAL_FILTER_ENABLED = True
        BILATERAL_FILTER_D = 22 #9 #4 #9  # Diâmetro da vizinhança do pixel
        BILATERAL_FILTER_SIGMA_COLOR = 75 #40 #75  # Filtra o espaço de cores sigma
        BILATERAL_FILTER_SIGMA_SPACE = 75 #40 #75  # Filtra o espaço coordenado sigma

        MEDIAN_BLUR_ENABLED = False
        MEDIAN_BLUR_KERNEL_SIZE = 3  # Tamanho do kernel para o desfoque mediano

        SHARPEN_FILTER_ENABLED = True
        SHARPEN_FILTER_KERNEL = np.array([[-1, -1, -1], [-1,  9, -1], [-1, -1, -1]])  # Kernel para o filtro de nitidez

        if GAUSSIAN_BLUR_ENABLED:
            img = cv2.GaussianBlur(img, GAUSSIAN_BLUR_KERNEL_SIZE, GAUSSIAN_BLUR_SIGMA_X, GAUSSIAN_BLUR_SIGMA_Y)

        if BILATERAL_FILTER_ENABLED:
            img = cv2.bilateralFilter(img, BILATERAL_FILTER_D, BILATERAL_FILTER_SIGMA_COLOR, BILATERAL_FILTER_SIGMA_SPACE)

        if MEDIAN_BLUR_ENABLED:
            img = cv2.medianBlur(img, MEDIAN_BLUR_KERNEL_SIZE)

        if SHARPEN_FILTER_ENABLED:
            img = cv2.filter2D(img, -1, SHARPEN_FILTER_KERNEL)

        return save_new_image(image_info, convert_cv2_to_pil(img))

    except FileNotFoundError:
        error = f"Arquivo não encontrado: {image_info.get('path')}"
    except UnidentifiedImageError:
        error = f"Arquivo não é uma imagem válida: {image_info.get('path')}"
    except OSError as e:
        error = f"Falha ao processar imagem '{image_info.get('path')}': {e}"
    except ValueError as e:
        error = f"Valor inválido ao salvar '{image_info.get('path')}': {e}"
    except KeyError as e:
        error = f"Formato não suportado para '{image_info.get('path')}': {e}"
    except Exception as e:
        error = f"Erro inesperado com '{image_info.get('path')}': {e}"

    image_info['status'] = False
    image_info['error'] = error

    return image_info, []

def noise_images(images_info):
    new_images_info = []
    old_images_info = []
    error_images_info = []

    with Pool(processes=cpu_count()) as pool:
        noise_result = pool.map(remove_noise_from_image, images_info)

    for new_image_info, old_image_info in noise_result:
        if new_image_info.get('status'):
            new_images_info.append(new_image_info)
            if isinstance(old_image_info, dict):
                old_images_info.append(old_image_info)
        else:
            error_images_info.append(new_image_info)

    return new_images_info, old_images_info, error_images_info

# --action from_grid --rows 3 --cols 3

# /Users/ro7rinke/Library/CloudStorage/GoogleDrive-ro7rinke2@gmail.com/My Drive/BoardGame/Brass Birmingham/final-print/cards/cards_grid_1.jpg

# --output_directory_path "/Users/ro7rinke/Documents/rr_image_utils/data/temp"

# --action to_pdf --file_name brass --output_directory_path "/Users/ro7rinke/Documents/rr_image_utils/data/temp"
    













# def test():
#     image_info = {
#         'external_source_path': '', #Path
#         'external_destination_path': '', #Path
#         'path': '', #Path
#         'session_id': '',
#         'id': '',
#         'old_id': '',
#         'original_name': '',
#         'original_extension': '',
#     }
#     return